import sys
import docx


# ARABIC VALUE RANGE
UNI_ARABIC_START = 1548
UNI_ARABIC_END = 1790

# DIGIT NUMBER RANGE
UNI_NUMBER_START = 48
UNI_NUMBER_END = 57

# SPECIFIC INCLUSIONS
UNI_SPACE_VAL = 32
UNI_QUESTION_MARK = 63
UNI_EXCLAMATION_POINT = 33
UNI_LEFT_PARENTHESIS = 40
UNI_RIGHT_PARENTHESIS = 41
UNI_QUOTATION_MARK = 34
UNI_DOLLAR_SIGN = 36
UNI_EURO_SIGN = 8364
UNI_COLON = 58
UNI_SEMICOLON = 59
UNI_AT_SIGN = 64
UNI_PERCENTAGE_SIGN = 37


def shouldKeepCharacter(character):
    charOrd = ord(character)
    if charOrd >= UNI_ARABIC_START and charOrd <= 1790 or  \
        charOrd >= UNI_NUMBER_START and charOrd <= UNI_NUMBER_END or \
        charOrd == UNI_SPACE_VAL or \
        charOrd == UNI_QUESTION_MARK or \
        charOrd == UNI_EXCLAMATION_POINT or \
        charOrd == UNI_LEFT_PARENTHESIS or \
        charOrd == UNI_RIGHT_PARENTHESIS or \
        charOrd == UNI_QUOTATION_MARK or \
        charOrd == UNI_QUESTION_MARK or \
        charOrd == UNI_DOLLAR_SIGN or \
        charOrd == UNI_EURO_SIGN or \
        charOrd == UNI_COLON or \
        charOrd == UNI_SEMICOLON or \
        charOrd == UNI_AT_SIGN or \
        charOrd == UNI_PERCENTAGE_SIGN:
            return True
    else:
        return False

# get filepath
filepath = str(sys.argv[1])
print("filepath used: ", filepath)

# get document
dirty_doc = docx.Document(filepath)
# print("document is " + str(len(dirty_doc.paragraphs)) + " paragraphs long")

# DEBUG LOGS: print paragraphs for viewing
# print("PRE-CLEANING")
# i=0
# while i<len(dirty_doc.paragraphs):
#     print("Paragraph ", str(i))
#     print(dirty_doc.paragraphs[i].text)
#     i+=1

# encode paragraphs to unicode


amended_paragraphs = []
for para in dirty_doc.paragraphs:
    newParagraph = ''
    for char in para.text:
        if shouldKeepCharacter(char):
            newParagraph += char
    amended_paragraphs.append(newParagraph)

print('cleaned paragraphs')
    
newDoc = docx.Document()

for para in amended_paragraphs:
    newDoc.add_paragraph(para)
newDoc.save('/mnt/c/Users/jmstr/OneDrive/Desktop/NEW_WORD.docx')


